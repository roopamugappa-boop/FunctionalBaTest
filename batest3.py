# =========================================================
# COMPLETE SINGLE FILE FLASK APP
# FUNCTIONAL BA ASSESSMENT WITH:
# - LOGIN
# - TIMER
# - MCQ
# - BRD WORD EDITOR
# - WORKFLOW BUILDER
# - WORKFLOW DRAWING CANVAS
# - GAP ANALYSIS
# - REVIEW
# - SQLITE
# =========================================================

from flask import (
    Flask,
    render_template_string,
    request,
    redirect,
    session,
    send_file,
    flash,
    url_for
)

from flask_session import Session

import sqlite3
from datetime import datetime

# =========================================================
# APP
# =========================================================

app = Flask(__name__)

app.secret_key = "secret_key"

app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_PERMANENT'] = False

Session(app)

DB = "assessment.db"

# =========================================================
# DATABASE
# =========================================================

def init_db():
    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS candidates(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            email TEXT,
            phone TEXT,
            start_time TEXT,
            submit_time TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS answers(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            candidate_id INTEGER,
            question TEXT,
            answer TEXT
        )
    ''')
    conn.commit()
    conn.close()

init_db()

# =========================================================
# BASE TEMPLATE (Improved for UI and Timer)
# =========================================================

BASE_HTML = '''
<!DOCTYPE html>
<html lang="en">
<head>
<title>Functional BA Assessment</title>

<link href="https://cdn.quilljs.com/1.3.6/quill.snow.css" rel="stylesheet">
<script src="https://cdn.quilljs.com/1.3.6/quill.js"></script>
<script src="https://cdn.jsdelivr.net/npm/sortablejs@latest/Sortable.min.js"></script>

<meta name="viewport" content="width=device-width, initial-scale=1.0">

<style>
:root {
    --accent: #2563eb;
    --danger: #e11d48;
    --success: #059669;
    --gray-bg: #eef2f7;
}
*{
    margin:0;
    padding:0;
    box-sizing:border-box;
    font-family:'Segoe UI', Arial, sans-serif;
}
body{
    background:var(--gray-bg);
    min-height: 100vh;
}
.navbar{
    background:#111827;
    color:white;
    padding:22px 10px;
    text-align:center;
    font-size:29px;
    font-weight:700;
    letter-spacing:0.1em;
    box-shadow:0 2px 12px #0002;
}
.container{
    width:97vw;
    max-width:1330px;
    margin:auto;
    margin-top:34px;
    margin-bottom:60px;
    background:white;
    border-radius:18px;
    padding:34px;
    box-shadow:0px 4px 32px rgba(0,0,0,0.11);
    min-height: 540px;
}
.question-box{
    background:linear-gradient(90deg,#f9fafb,#e9eefe 85%);
    padding:24px;
    border-radius:14px;
    margin-bottom:28px;
    border-left:7px solid var(--accent);
    box-shadow:0 1px 16px #2d6be30b;
    transition: box-shadow 0.12s;
}
.question{
    font-size:20px;
    font-weight:600;
    margin-bottom:14px;
    color: #10223a;
    letter-spacing: 0.01em;
}
.option{
    background:white;
    border:1.6px solid #d0d3e2;
    padding:14px 17px;
    border-radius:12px;
    margin-bottom:10px;
    transition: box-shadow 0.18s;
    font-size:15.7px;
}
.option input[type=radio] {
    accent-color: var(--accent);
    margin-right: 8px;
    vertical-align: middle;
}
textarea, input[type=text]{
    width:100%;
    padding:16px;
    border-radius:11px;
    border:1.6px solid #c0c4d7;
    margin-bottom:22px;
    transition: border-color 0.18s;
    font-size:17px;
}
textarea:focus, input[type=text]:focus {
    outline:none;
    border-color:var(--accent);
    background: #f5f9ff;
}
.btn{
    background:linear-gradient(90deg,#2563eb,#1fb6ff 80%);
    color:white;
    border:none;
    padding:15px 29px;
    border-radius:10px;
    font-size:18px;
    cursor:pointer;
    font-weight: 600;
    box-shadow:0 2px 8px #2563eb18;
    letter-spacing: .01em;
    transition: background 0.17s, box-shadow 0.19s, transform 0.09s;
}
.btn:active{ transform:scale(0.965);}
.btn:hover{
    background: linear-gradient(90deg,#1d4ed8,#2563eb 60%);
}

.timer{
    position:fixed;
    top:20px;
    right:20px;
    background-color:white;
    border:2.5px solid var(--accent);
    color:#22223b;
    padding:11px 24px;
    border-radius:12px;
    font-size:28px;
    z-index:1019;
    box-shadow: 0 4px 32px #2563eb15;
    font-family:'Roboto Mono',monospace;
    font-weight: 700;
    display: flex;
    align-items: center;
    gap:12px;
    min-width:150px;
    transition: background 0.18s, border 0.18s, color 0.2s;
}
.timer-warning {
    border-color: #fbbf24 !important;
    background: #fff9eb !important;
    color: #b45309 !important;
    animation: timerwarn 1s infinite alternate;
}
.timer-danger {
    border-color: var(--danger) !important;
    background: #fff3f6 !important;
    color: var(--danger) !important;
    animation: timerpulse .8s infinite alternate;
}
@keyframes timerwarn {
    0%{ box-shadow: 0 0 0 0 #f59e4211;}
    100%{ box-shadow: 0 0 16px 0 #fbbf2445;}
}
@keyframes timerpulse {
    0%{ box-shadow: 0 0 0 0 #e11d4810;}
    100%{ box-shadow: 0 0 18px 0 #e11d4833;}
}
.progress{
    width:100%;
    height:17px;
    background:#e5e7ee;
    border-radius:20px;
    overflow:hidden;
    margin-bottom:29px;
    border:1.2px solid #c5cdf3;
}
.progress-bar{
    height:100%;
    background: linear-gradient(90deg,#2563eb 40%,#38a3f4 95%);
    transition: width .5s cubic-bezier(.1,2,.7,.98);
}
.workflow-step{
    background: #e0e8fb;
    padding:16px;
    border-radius:12px;
    margin-bottom:11px;
    cursor:move;
    font-weight:600;
    font-size:16px;
    border: 1.1px solid #b1c6ff;
    box-shadow: 0 2px 10px #2563eb08;
}
#editor{
    height:420px;
    background:white;
    border-radius:10px;
    border:1px solid #c0c4d7;
}
.nav-links{
    display:flex;
    gap:12px;
    margin-bottom:28px;
    flex-wrap:wrap;
    justify-content: flex-start;
}
.nav-links a{
    background:#111827;
    color:white;
    padding:12px 20px;
    text-decoration:none;
    border-radius:9px;
    font-size: 17px;
    letter-spacing: .01em;
    transition: background .1s;
}
.nav-links a:hover, .nav-links a.active{
    background:var(--accent);
}
.card-option{
    padding:16px;
    background:#f3f4f6;
    border-radius:12px;
    margin-bottom:13px;
    cursor:pointer;
    border:2.4px solid transparent;
    font-size: 16.5px;
    transition: border .16s, background .18s;
}
.card-option.selected{
    border-color:var(--success);
    background:#d1fae5;
}
::-webkit-scrollbar {width: 8px;}
::-webkit-scrollbar-thumb {background: #e9edf1; border-radius:4px;}
::-webkit-scrollbar-track {background:transparent;}
@media (max-width:600px){
    .container{padding:6vw;}
    .timer{ font-size:20px; padding: 9px 14px; min-width: 98px;}
    .navbar{font-size:21px; padding: 13px 4px;}
}
</style>

</head>

<body>

<div class="timer" id="timer" style="display: flex; align-items: center; justify-content: flex-end; background: #f8fafc; border: 1.5px solid #c7d2fe; border-radius: 9px; box-shadow: 0 2px 12px #0001; padding: 10px 28px; font-size: 22px; font-weight: 500; color: #253262; margin-top: 18px; margin-bottom: 20px; min-width: 180px;">
    <span id="timer-icon" style="font-size: 26px; vertical-align: middle; margin-right: 14px; color: #6366f1;">⏰</span>
    <span id="timer-text" style="font-family: 'Consolas', 'Menlo', monospace; font-size: 1.22em; letter-spacing: 0.04em; color: #21293a; background: #e0e7ff; padding: 7px 18px; border-radius: 7px; min-width: 80px; text-align: center; border: 1.1px solid #a5b4fc; box-shadow: 0 1px 6px #5364a522;">75:00</span>
</div>

<div class="navbar">
    <span style="color:red;font-weight:bold;">T</span>ime Line Investments Pvt Ltd | Functional BA Assessment Portal
</div>

<div class="container">

    <div class="progress">
        <div class="progress-bar" style="width:{{ progress }}%"></div>
    </div>
    <div class="nav-links">
        <a href="/mcq" {% if '/mcq' in request.path %}class="active"{% endif %}>MCQ</a>
        <a href="/brd" {% if '/brd' in request.path %}class="active"{% endif %}>BRD Editor</a>
        <a href="/workflow" {% if '/workflow' in request.path and not '/workflow_canvas' in request.path %}class="active"{% endif %}>Workflow</a>
        <a href="/workflow_canvas" {% if '/workflow_canvas' in request.path %}class="active"{% endif %}>Workflow Canvas</a>
        <a href="/gap" {% if '/gap' in request.path %}class="active"{% endif %}>KYC & Payment Gateway</a>
        <a href="/review" {% if '/review' in request.path %}class="active"{% endif %}>Review</a>
    </div>
    {% with messages = get_flashed_messages(with_categories=true) %}
      {% if messages %}
          <ul style="list-style:none; padding:0; margin-bottom:19px;">
          {% for category, msg in messages %}
            <li style="margin-bottom:6px;">
              <span style="display:inline-block;padding:11px 18px;border-radius:8px; font-size:15px; {% if category == 'error' %}background:#fff4f3; color:#ea1e49; border:1px solid #ffdada;{% elif category == 'warning' %}background:#fdf7c3; color:#b45309; border:1px solid #ffe29d;{% else %}background:#e7fbee; color:#059669; border:1.2px solid #8cfbea;{% endif %} box-shadow:0 2px 16px #0001;">
              <b>{% if category == 'error' %}❗{% elif category == 'warning' %}⚠️{% else %}✔{% endif %}</b>
              {{msg}}
              </span>
            </li>
          {% endfor %}
          </ul>
      {% endif %}
    {% endwith %}
    {{ body|safe }}
</div>

<script>
// Improved Timer: visual warnings, autosubmit, blinking at < 1 min, session timeout, etc.
(function(){
const timerEl = document.getElementById("timer");
const timerText = document.getElementById("timer-text");
const timerIcon = document.getElementById("timer-icon");
const total = 75*60; // seconds

// Use localStorage for resilience
if(!localStorage.getItem("timeLeft")){
    localStorage.setItem("timeLeft", total);
}
let timeLeft = parseInt(localStorage.getItem("timeLeft")) || total;

function updateDisplay() {
    let mins = Math.floor(timeLeft/60);
    let secs = timeLeft%60;
    timerText.textContent = String(mins).padStart(2,'0')+":"+String(secs).padStart(2,'0');
    timerEl.classList.remove('timer-warning','timer-danger');
    // Visual indicators:
    if (timeLeft <= 60) {
        timerEl.classList.add('timer-danger');
        timerIcon.textContent = "⏳";
        // Blink text when < 60s
        timerText.style.animation = "blinker .75s linear infinite";
    }
    else if (timeLeft <= 240) {
        timerEl.classList.add('timer-warning');
        timerIcon.textContent = "⌛";
        timerText.style.animation = "";
    }
    else {
        timerEl.classList.remove('timer-warning','timer-danger');
        timerIcon.textContent = "⏰";
        timerText.style.animation = "";
    }
}
updateDisplay();
window.onblur = function(){ // Save when lost focus!
    localStorage.setItem("timeLeft", timeLeft);
};

let interval = setInterval(function(){
    timeLeft = parseInt(localStorage.getItem("timeLeft")) || total;
    if(timeLeft <= 0){
        timerText.textContent = "00:00";
        timerEl.classList.add('timer-danger');
        localStorage.removeItem("timeLeft");
        setTimeout(function(){ window.location = "/submit"; }, 650);
        clearInterval(interval);
        return;
    }
    timeLeft--;
    localStorage.setItem("timeLeft", timeLeft);
    updateDisplay();
},1000);

// Animate blinking warning at <1 minute
var style = document.createElement('style');
style.innerHTML = `@keyframes blinker { 50% { opacity: 0.4; }}`;
document.head.appendChild(style);

})();
</script>
</body>
</html>
'''

# =========================================================
# LOGIN
# =========================================================

@app.route('/', methods=['GET','POST'])
def login():
    if request.method == 'POST':
        name = request.form['name'].strip()
        email = request.form['email'].strip()
        phone = request.form['phone'].strip()
        if not (name and email and phone):
            flash("Please enter all fields.", "warning")
        else:
            conn = sqlite3.connect(DB)
            c = conn.cursor()
            c.execute('''
                INSERT INTO candidates(
                    name, email, phone, start_time
                )
                VALUES(?,?,?,?)
            ''', (
                name,
                email,
                phone,
                str(datetime.now())
            ))
            conn.commit()
            session['candidate_id'] = c.lastrowid
            conn.close()
            localstorage_js = '''
            <script> localStorage.setItem("timeLeft", 75*60); </script>
            '''
            # Reset timer when login
            return '''
            <script>
                localStorage.setItem("timeLeft", 75*60);
                window.location.href = "/mcq";
            </script>
            <div style="font-size:1.8em;text-align:center;margin-top:90px">
                <a href="/mcq" style="color:#2653a8;text-decoration:underline;font-weight:600;">
                    👉 Click here to start your assessment
                </a>
            </div>
            '''
       
    body = '''
    <h2 style="margin-bottom:12px;">Candidate Login</h2>
    <div class="question-box" style="max-width:480px;margin:auto;">
    <form method="POST" autocomplete="off">
        <input type="text"
               name="name"
               placeholder="Candidate Name"
               required
               style="margin-bottom:14px">
        <input type="text"
               name="email"
               placeholder="Email"
               required
               style="margin-bottom:14px">
        <input type="text"
               name="phone"
               placeholder="Phone"
               required
               pattern="[0-9+\\- ()]{8,20}"
               title="Enter a valid phone number"
               style="margin-bottom:22px">
        <button class="btn" style="width:100%;font-size:19px;">
        🚀 Start Assessment
        </button>
    </form>
    </div>
    '''
    return render_template_string(
        BASE_HTML,
        body=body,
        progress=5
    )

# =========================================================
# MCQ
# =========================================================

@app.route('/mcq', methods=['GET','POST'])
def mcq():
    if request.method == 'POST':
        conn = sqlite3.connect(DB)
        c = conn.cursor()
        # Insert all questions and answers into the db
        questions = [
            ('Purpose of KYC', request.form.get('q1')),
            ('UIDAI Authority', request.form.get('q2')),
            ('Purpose of OTP-based Aadhaar authentication', request.form.get('q3')),
            ('System action for not received OTP after correct Aadhaar', request.form.get('q4')),
            ('Purpose of Virtual ID (VID)', request.form.get('q5')),
            ('What should happen after 3 failed OTP attempts?', request.form.get('q6')),
            ('Importance of audit logs in KYC systems', request.form.get('q7')),
            ('First stage of payment transaction lifecycle', request.form.get('q8')),
            # Add further MCQ/short questions here as needed
        ]
        answers = [
            (
                session['candidate_id'],
                q,
                a
            )
            for (q, a) in questions
        ]
        c.executemany('''
            INSERT INTO answers(
                candidate_id,
                question,
                answer
            )
            VALUES(?,?,?)
        ''', answers)
        conn.commit()
        conn.close()
        return redirect('/brd')
    body = '''
    <h2 style="margin-bottom:14px;">KYC & Aadhaar MCQ + Payments Systems</h2>
    <form method="POST">

    <!-- Q1 -->
    <div class="question-box">
        <div class="question">
            Q1. What is the primary purpose of KYC in fintech systems?
        </div>
        <div class="option">
            <input type="radio" name="q1" value="A" required> A) Marketing customers
        </div>
        <div class="option">
            <input type="radio" name="q1" value="B"> B) Customer verification and fraud prevention
        </div>
        <div class="option">
            <input type="radio" name="q1" value="C"> C) UI enhancement
        </div>
        <div class="option">
            <input type="radio" name="q1" value="D"> D) Employee onboarding
        </div>
    </div>

    <!-- Q2 -->
    <div class="question-box">
        <div class="question">
            Q2. Which authority manages Aadhaar authentication?
        </div>
        <div class="option">
            <input type="radio" name="q2" value="A" required> A) RBI
        </div>
        <div class="option">
            <input type="radio" name="q2" value="B"> B) NPCI
        </div>
        <div class="option">
            <input type="radio" name="q2" value="C"> C) UIDAI
        </div>
        <div class="option">
            <input type="radio" name="q2" value="D"> D) SEBI
        </div>
    </div>

    <!-- Q3 -->
    <div class="question-box">
        <div class="question">
            Q3. What is the purpose of OTP-based Aadhaar authentication?
        </div>
        <div class="option">
            <input type="radio" name="q3" value="A" required> A) Tax filing
        </div>
        <div class="option">
            <input type="radio" name="q3" value="B"> B) Identity verification
        </div>
        <div class="option">
            <input type="radio" name="q3" value="C"> C) Salary processing
        </div>
        <div class="option">
            <input type="radio" name="q3" value="D"> D) Payment settlement
        </div>
    </div>

    <!-- Q4 -->
    <div class="question-box">
        <div class="question">
            Q4. Customer entered Aadhaar correctly but OTP is not received. What should system do <b>first</b>?
        </div>
        <div class="option">
            <input type="radio" name="q4" value="A" required> A) Approve manually
        </div>
        <div class="option">
            <input type="radio" name="q4" value="B"> B) Retry OTP generation
        </div>
        <div class="option">
            <input type="radio" name="q4" value="C"> C) Skip KYC
        </div>
        <div class="option">
            <input type="radio" name="q4" value="D"> D) Block customer
        </div>
    </div>

    <!-- Q5 -->
    <div class="question-box">
        <div class="question">
            Q5. What is Virtual ID (VID) used for?
        </div>
        <div class="option">
            <input type="radio" name="q5" value="A" required> A) Payment authorization
        </div>
        <div class="option">
            <input type="radio" name="q5" value="B"> B) Replacing Aadhaar number for privacy
        </div>
        <div class="option">
            <input type="radio" name="q5" value="C"> C) Settlement process
        </div>
        <div class="option">
            <input type="radio" name="q5" value="D"> D) Bank account verification
        </div>
    </div>

    <!-- Q6 -->
    <div class="question-box">
        <div class="question">
            Q6. What should happen after 3 failed OTP attempts?<br>
            <small>(ONE LINE ANSWER)</small>
        </div>
        <input type="text" name="q6" class="input" style="width:99%;margin-top:10px;" placeholder="E.g. Temporary block / retry cooldown / security validation">
    </div>

    <!-- Q7 -->
    <div class="question-box">
        <div class="question">
            Q7. What is the importance of audit logs in KYC systems?<br>
            <small>(SHORT ANSWER) E.g. Tracking activities for compliance, security, monitoring, and investigations.</small>
        </div>
        <textarea name="q7" class="input" rows="2" style="width:99%;margin-top:10px;" placeholder="Your answer"></textarea>
    </div>

    

    <!-- Q8 -->
    <div class="question-box">
        <div class="question">
            Q8. What is the first stage of payment transaction lifecycle?
        </div>
        <div class="option">
            <input type="radio" name="q8" value="A" required> A) Settlement
        </div>
        <div class="option">
            <input type="radio" name="q8" value="B"> B) Reconciliation
        </div>
        <div class="option">
            <input type="radio" name="q8" value="C"> C) Authorization
        </div>
        <div class="option">
            <input type="radio" name="q8" value="D"> D) Initiation
        </div>
    </div>

    <!-- Add additional payment questions as per requirements -->

    <button class="btn" style="float:right;margin-top:14px;">
        Save & Next →
    </button>
    </form>
    '''
    return render_template_string(
        BASE_HTML,
        body=body,
        progress=20
    )

# =========================================================
# BRD EDITOR
# =========================================================

@app.route('/brd', methods=['GET', 'POST'])
def brd():
    import re
    import io
    import base64
    from flask import send_file, request, session, redirect, render_template_string
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup

    # Helper to add header/footer to each section
    def add_header_footer(doc, header_text="BRD / FRD Document", footer_text="Confidential – Company Name"):
        for section in doc.sections:
            # Header
            header = section.header
            header.is_linked_to_previous = False
            if header.paragraphs:
                p = header.paragraphs[0]
                p.text = header_text
            else:
                p = header.add_paragraph(header_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # More robust font application
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(12)

            # Footer
            footer = section.footer
            footer.is_linked_to_previous = False
            if footer.paragraphs:
                p2 = footer.paragraphs[0]
                p2.text = footer_text + " | Page "
            else:
                p2 = footer.add_paragraph(footer_text + " | Page ")
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p2.runs:
                run.font.size = Pt(10)
            # Add placeholder for page number
            run = p2.add_run(" { PAGE } ")
            run.font.size = Pt(10)

    # Enhanced HTML->docx conversion
    def html_to_docx(html, doc):
        soup = BeautifulSoup(html or "", "html.parser")

        def parse_elem(elem, parent=None):
            if elem.name is None:  # NavigableText
                text = elem.string
                if text and text.strip():
                    text = text.strip('\n')
                    if parent and parent.name in ['b', 'strong']:
                        p = doc.add_paragraph()
                        run = p.add_run(text)
                        run.bold = True
                    elif parent and parent.name in ['i', 'em']:
                        p = doc.add_paragraph()
                        run = p.add_run(text)
                        run.italic = True
                    elif parent and parent.name == 'u':
                        p = doc.add_paragraph()
                        run = p.add_run(text)
                        run.underline = True
                    else:
                        doc.add_paragraph(text)
                return
            tag = elem.name.lower()
            if tag in ['h1','h2','h3','h4','h5','h6']:
                lv = int(tag[1])
                doc.add_heading(elem.get_text(strip=True), level=lv)
            elif tag == 'ul':
                for li in elem.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
            elif tag == 'ol':
                for li in elem.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(strip=True), style='List Number')
            elif tag == 'table':
                rows = elem.find_all('tr')
                if rows:
                    cols = rows[0].find_all(['td', 'th'])
                    table = doc.add_table(rows=len(rows), cols=len(cols))
                    table.style = 'Table Grid'
                    for i, tr in enumerate(rows):
                        for j, td in enumerate(tr.find_all(['td', 'th'])):
                            cell = table.cell(i, j)
                            # Slightly improved: retain bold for th
                            is_header = td.name == 'th'
                            txt = td.get_text(strip=True)
                            if is_header:
                                run = cell.paragraphs[0].add_run(txt)
                                run.bold = True
                            else:
                                cell.text = txt
            elif tag == 'img':
                src = elem.get('src', '')
                if src.startswith('data:image/') and "base64," in src:
                    match = re.match(r'data:image/(.*?);base64,(.*)', src)
                    if match:
                        ext, b64 = match.groups()
                        try:
                            image_stream = io.BytesIO(base64.b64decode(b64))
                            doc.add_picture(image_stream, width=Inches(4))
                        except Exception as ex:
                            doc.add_paragraph(f"[[Image could not be embedded: {str(ex)}]]")
            elif tag == 'a':
                href = elem.get('href', '')
                # Visually style as a link (blue, underline, smaller font)
                p = doc.add_paragraph()
                run = p.add_run(elem.get_text(strip=True))
                run.underline = True
                run.font.size = Pt(10)
                p.add_run(f" ({href})").italic = True
            elif tag == 'blockquote':
                doc.add_paragraph(elem.get_text(strip=True), style='Intense Quote')
            elif tag == 'code':
                doc.add_paragraph(elem.get_text(strip=True), style='Intense Quote')
            elif tag == 'pre':
                doc.add_paragraph(elem.get_text(strip=False), style='Intense Quote')
            elif tag == 'hr':
                doc.add_page_break()
            elif tag in ['b', 'strong', 'i', 'em', 'u']:
                # Styling handled above for NavigableText
                for c in elem.children:
                    parse_elem(c, elem)
                return
            # Recurse for other tags and children
            for c in elem.children:
                parse_elem(c, elem)

        parse_elem(soup.body if soup.body else soup)
        add_header_footer(doc)
        return doc

    if request.method == 'POST':
        brd_content = request.form.get('brd_content')
        export_word = request.form.get('export_word')
        if export_word == "1":
            doc = Document()
            doc.add_heading('BRD / FRD Document', 0)
            html_to_docx(brd_content, doc)
            memfile = io.BytesIO()
            doc.save(memfile)
            memfile.seek(0)
            return send_file(memfile, download_name="BRD_FRD.docx", as_attachment=True)
        else:
            import sqlite3
            conn = sqlite3.connect(DB)
            c = conn.cursor()
            c.execute(
                '''
                INSERT INTO answers(candidate_id, question, answer)
                VALUES (?, ?, ?)
                ''',
                (session['candidate_id'], 'BRD Editor', brd_content)
            )
            conn.commit()
            conn.close()
            return redirect('/workflow')

    body = '''
    <h2>BRD Text Editor</h2>
    <form method="POST" onsubmit="saveEditor()" id="brd-form" style="margin-bottom:0;">
        <div style="margin-bottom:14px;font-size:1.01em;color:#555;">
            <div style="font-size:1.07em; line-height:1.62;">
                <b style="font-size:1.13em;color:#253262;">Professional Instructions:</b>
                <br>
                You are tasked with drafting a <b>Business Requirements Document (BRD)</b> focused on the <b>integration of a Payment Gateway</b> within a fintech environment.
                <br><br>
                <span style="color:#444;">
                Your BRD should be precise, well-structured, and tailored for a professional audience. Please ensure your document conveys a clear understanding of the project scope, objectives, and all associated requirements. Pay careful attention to both business and technical aspects.
                </span>
                <br><br>
                <ol style="margin-left:30px; padding-left:4px;">
                    <li>
                        <b>Business Objectives & Scope:</b> Clearly outline the project objectives and the boundaries of the integration initiative.
                    </li>
                    <li>
                        <b>Stakeholders:</b> Identify the relevant stakeholders and their roles in the project.
                    </li>
                    <li>
                        <b>Functional Requirements:</b> Detail all core and supporting functionalities required for the payment gateway integration.
                    </li>
                    <li>
                        <b>Non-functional Requirements:</b> Specify key performance metrics, security standards, usability needs, scalability, and compliance obligations.
                    </li>
                    <li>
                        <b>User Stories / Use Cases:</b> Describe relevant workflows from an end-user perspective to illustrate functional requirements.
                    </li>
                    <li>
                        <b>Process Flows & Workflow Diagrams:</b> Provide clear diagrams or step-by-step lists to visually communicate the integration process.
                    </li>
                    <li>
                        <b>Integration Points:</b> Summarize external/internal system interfaces, data exchanges, and validation or security checkpoints as appropriate.
                    </li>
                    <li>
                        <b>Acceptance Criteria:</b> List measurable criteria to evaluate if requirements have been met.
                    </li>
                    <li>
                        <b>Formatting & Visuals:</b> Utilize headings, tables, bullet/numbered lists, and diagrams or images to enhance clarity and professionalism.
                    </li>
                    <li>
                        <b>Focus:</b> You may choose either <b>Payment Gateway</b> or <b>KYC process</b> for your BRD—ensure the document is complete, structured, and easy to navigate.
                    </li>
                </ol>
                <div style="color: #6366f1; font-size:0.99em; margin-top:10px;">
                    <b>Tip:</b> Leverage formatting, tables, and visuals throughout your document for greater impact and clarity.
                </div>
            </div>
       
 
            
            <b>Note:</b> Exported Word retains most formatting & images and adds header/footer to every page.
        </div>
        <div id="editor"></div>
        <input type="hidden" name="brd_content" id="brd_content">
        <input type="hidden" name="export_word" id="export_word" value="0">
        <br>
        <button class="btn" style="margin-top:12px;float:right;">Save & Next →</button>
        <button class="btn" type="button" style="margin-top:12px;margin-right:15px;background:linear-gradient(90deg,#059669,#0ea5e9 85%);" onclick="exportWord()">⬇ Export as Word (with header/footer)</button>
    </form>
    <script src="https://cdn.quilljs.com/1.3.7/quill.js"></script>
    <link href="https://cdn.quilljs.com/1.3.7/quill.snow.css" rel="stylesheet">
    <script>
    var quill = new Quill('#editor', {
        theme: 'snow',
        modules: {
            toolbar: [
                [{ 'header': [1, 2, 3, 4, 5, 6, false] }],
                [{ 'font': [] }],
                [{ 'size': ['small', false, 'large', 'huge'] }],
                ['bold', 'italic', 'underline', 'strike'],
                [{ 'color': [] }, { 'background': [] }],
                [{ 'script': 'sub'}, { 'script': 'super' }],
                [{ 'list': 'ordered' }, { 'list': 'bullet'}, { 'list': 'check'}],
                [{ 'indent': '-1' }, { 'indent': '+1' }],
                [{ 'direction': 'rtl' }], [{ 'align': [] }],
                ['blockquote', 'code-block'],
                ['link', 'image', 'video', 'formula'],
                ['clean']
            ]
        }
    });

    // Editor data handling for submit/export
    function saveEditor() {
        document.getElementById('brd_content').value = quill.root.innerHTML;
        document.getElementById('export_word').value = "0";
    }
    function exportWord() {
        document.getElementById('brd_content').value = quill.root.innerHTML;
        document.getElementById('export_word').value = "1";
        document.getElementById('brd-form').submit();
    }

    // User-friendly image upload to Quill (Base64)
    quill.getModule("toolbar").addHandler('image', function() {
      var input = document.createElement('input');
      input.type = 'file';
      input.accept = 'image/*';
      input.click();
      input.onchange = function() {
        var file = input.files[0];
        if (!file) return;
        var reader = new FileReader();
        reader.onload = function(e) {
          var range = quill.getSelection(true) || { index: quill.getLength() };
          quill.insertEmbed(range.index, 'image', e.target.result, 'user');
        };
        reader.readAsDataURL(file);
      };
    });

    // Optional: auto-initialize editor if returning to edit
    if (window.localStorage.getItem('brd_content')) {
        quill.root.innerHTML = window.localStorage.getItem('brd_content');
    }
    // Save draft in localStorage to prevent browser loss
    quill.on('text-change', function() {
        window.localStorage.setItem('brd_content', quill.root.innerHTML);
    });
    </script>
    <style>
        #editor {
            background: #fff;
            min-height: 350px;
            border-radius: 8px;
            border: 1.5px solid #dfe6eb;
            box-shadow: 0px 2px 12px #0001;
        }
    </style>
    '''
    return render_template_string(
        BASE_HTML,
        body=body,
        progress=40
    )

# =========================================================
# WORKFLOW BUILDER
# =========================================================

@app.route('/workflow', methods=['GET','POST'])
def workflow():
    if request.method == 'POST':
        workflow = request.form.get('workflow_order')
        conn = sqlite3.connect(DB)
        c = conn.cursor()
        c.execute('''
            INSERT INTO answers(
                candidate_id,
                question,
                answer
            )
            VALUES(?,?,?)
        ''', (
            session['candidate_id'],
            'Workflow Order',
            workflow
        ))
        conn.commit()
        conn.close()
        return redirect('/workflow_canvas')
    body = '''
    <h2>Payment Gateway Integration Workflow Builder</h2>
    <form method="POST">
    <div id="workflow">
        <div class="workflow-step">
            Requirement Gathering<br>
            <small style="color:#64748b;">Define payment flow, failure flow, refund, retry<br>➜ Deliverable: <b>FRS</b></small>
        </div>
        <div class="workflow-step">
            Identify Business Need<br>
            <small style="color:#64748b;">Understand why payment gateway is required and payment methods (UPI, cards, etc.)<br>➜ Deliverable: <b>BRD</b></small>
        </div>
        <div class="workflow-step">
            Process Mapping<br>
            <small style="color:#64748b;">Create end-to-end flow diagram<br>➜ Deliverable: <b>Process Flow</b></small>
        </div>
        <div class="workflow-step">
            Vendor Evaluation<br>
            <small style="color:#64748b;">Compare vendors based on pricing, APIs, SLA, security<br>➜ Deliverable: <b>Vendor Comparison Matrix</b></small>
        </div>
        <div class="workflow-step">
            Technical Planning<br>
            <small style="color:#64748b;">Define APIs, webhooks, security, DB changes<br>➜ Deliverable: <b>SDD</b></small>
        </div>
        <div class="workflow-step">
            Commercial & Legal Approval<br>
            <small style="color:#64748b;">Finalize contract, NDA, SLA<br>➜ Deliverable: <b>Signed Agreement</b></small>
        </div>
        <div class="workflow-step">
            API Integration (Development)<br>
            <small style="color:#64748b;">Develop payment, status, refund APIs</small>
        </div>
        <div class="workflow-step">
            Sandbox Setup<br>
            <small style="color:#64748b;">Get test credentials from vendor<br>➜ Deliverable: <b>Sandbox Access</b></small>
        </div>
        <div class="workflow-step">
            SIT Testing<br>
            <small style="color:#64748b;">Test system integration and API responses</small>
        </div>
        <div class="workflow-step">
            Testing Preparation<br>
            <small style="color:#64748b;">Create test cases (success, failure, edge cases)<br>➜ Deliverable: <b>Test Cases, RTM</b></small>
        </div>
        <div class="workflow-step">
            UAT Testing<br>
            <small style="color:#64748b;">Business validates end-to-end flow<br>➜ Deliverable: <b>UAT Sign-Off</b></small>
        </div>
        <div class="workflow-step">
            Go-Live Readiness<br>
            <small style="color:#64748b;">Check SSL, monitoring, rollback plan<br>➜ Deliverable: <b>Go-Live Approval</b></small>
        </div>
        <div class="workflow-step">
            Production Setup<br>
            <small style="color:#64748b;">Get live API keys and configure system</small>
        </div>
        <div class="workflow-step">
            Post Go-Live Monitoring<br>
            <small style="color:#64748b;">Monitor failures, refunds, settlements<br>➜ Deliverable: <b>Hypercare Report</b></small>
        </div>
        <div class="workflow-step">
            Test Transaction (Go-Live)<br>
            <small style="color:#64748b;">Do ₹1 transaction and verify all flows<br>➜ Deliverable: <b>Production Validation Report</b></small>
        </div>
    </div>
    <input type="hidden" name="workflow_order" id="workflow_order">
    <br>
    <button class="btn" onclick="saveWorkflow()" style="margin-top:10px;float:right;">
        Save & Next →
    </button>
    </form>
    <script>
    new Sortable(
        document.getElementById('workflow'),
        { animation: 155 }
    );
    function saveWorkflow(){
        let steps = document.querySelectorAll('.workflow-step');
        let arr = [];
        steps.forEach(function(s){ arr.push(s.innerText.trim()); });
        document.getElementById('workflow_order').value = arr.join(' -> ')
    }
    </script>
    '''
    return render_template_string(
        BASE_HTML,
        body=body,
        progress=60
    )

# =========================================================
# WORKFLOW CANVAS
@app.route('/workflow_canvas', methods=['GET', 'POST'])
def workflow_canvas():
    import json
    import sqlite3
    from flask import render_template_string, request, redirect, session, flash, url_for

    if request.method == 'POST':
        canvas_data = request.form.get('canvas_data', '')
        candidate_id = session.get('candidate_id')
        if not candidate_id:
            flash("Your session expired. Please sign in again to continue.", "error")
            return redirect(url_for('workflow_canvas'))
        if not canvas_data or canvas_data.strip() == "[]":
            flash("🎨 It looks like the canvas is still empty. Please sketch your workflow before saving!", "warning")
            return redirect(url_for('workflow_canvas'))
        try:
            conn = sqlite3.connect(DB)
            c = conn.cursor()
            c.execute("""
                INSERT INTO answers(candidate_id, question, answer)
                VALUES (?, ?, ?)
            """, (candidate_id, "Workflow Drawing", canvas_data))
            conn.commit()
        except Exception as e:
            flash("Oops! Something went wrong while saving your workflow. Please try again. " + str(e), "error")
            return redirect(url_for('workflow_canvas'))
        finally:
            conn.close()
        flash("✅ Workflow saved successfully! Redirecting to next step ...", "success")
        return redirect('/gap')

    body = '''
    <h2 style="margin-bottom: 12px;">Draw Your KYC Workflow</h2>
    <div class="question-box" style="padding-bottom:24px;">
      <div style="background: #e9f6ff; border-radius: 10px; margin-bottom: 18px; padding: 14px 22px; color: #2c334e;">
        <span style="font-size:18px;font-weight:500;">🖌️ &nbsp;Visualize your KYC workflow step by step.</span><br>
        <span style="color:#6b7280;">Pick a tool below, then click or drag on the canvas. You can move or edit with <b>Move/Select</b>, undo mistakes, and try as many times as you like.<br>
        <i>Tip: Keyboard shortcuts like Ctrl+Z (Undo), Ctrl+Y (Redo), and <b>P</b> (Pen), <b>V</b> (Move/Select) make it even faster!</i></span>
      </div>
      <div style="display:flex;gap:10px;flex-wrap:wrap;margin-bottom:16px;align-items:center;">
        <button type="button" class="btn" onclick="setTool('pen')" id="tool-pen">✏️ Pen</button>
        <button type="button" class="btn" onclick="setTool('rect')" id="tool-rect">⬛ Rectangle</button>
        <button type="button" class="btn" onclick="setTool('circle')" id="tool-circle">⚪ Circle</button>
        <button type="button" class="btn" onclick="setTool('line')" id="tool-line">📏 Line</button>
        <button type="button" class="btn" onclick="setTool('arrow')" id="tool-arrow">➡️ Arrow</button>
        <button type="button" class="btn" onclick="setTool('connector')" id="tool-connector">🔗 Connect</button>
        <button type="button" class="btn" onclick="setTool('text')" id="tool-text">🔤 Text</button>
        <button type="button" class="btn" onclick="setTool('select')" id="tool-select">🤚 Move/Select</button>
        <button type="button" class="btn" onclick="undo()" title="Undo (Ctrl+Z)">↩️ Undo</button>
        <button type="button" class="btn" onclick="redo()" title="Redo (Ctrl+Y)">↪️ Redo</button>
        <button type="button" class="btn" onclick="clearCanvas()">🗑️ Clear</button>
        <span style="margin-left:14px;font-size:15px;">Color:</span>
        <input type="color" id="colorPicker" value="#2563eb" title="Select color" style="margin-left:4px;">
        <span style="margin-left:12px;font-size:15px;">Brush:</span>
        <input type="range" id="brushSize" min="1" max="20" value="3" style="width:70px;vertical-align:middle;" title="Brush size">
      </div>
      <canvas id="canvas" width="1100" height="650"
        style="border:2.5px solid #2563eb;width:100%;background:white;box-shadow:0 2px 16px #0001;border-radius:11px;transition:border-color 0.18s;">
      </canvas>
      <form method="POST" id="canvasForm" style="margin-top:22px;text-align:right;">
        <input type="hidden" name="canvas_data" id="canvas_data">
        <button type="button" class="btn" onclick="saveCanvas()" id="save-btn" style="font-weight:700;font-size:17px;">
          💾 Save & Continue
        </button>
      </form>
      <div id="canvas-message" style="margin-top:18px;color:#2563eb;font-size:15px;display:none;transition:opacity 0.2s;"></div>
      <div style="font-size:14px; margin-top:16px; color:#333;">
        <b>Connect Mode:</b> Select "🔗 Connect", then click two objects to draw a proper connecting line from one object to the next object.
      </div>
    </div>

<script>
// ... original canvas editor JS exactly as before ...
const canvas = document.getElementById("canvas");
const ctx = canvas.getContext("2d");

let tool = "pen";
let actions = [];
let redoStack = [];
let toolButtons = ["pen", "rect", "circle", "line", "arrow", "connector", "text", "select"];

let drawing = false, startX = 0, startY = 0, currentPath = [];
let selected = null, dragging = false, dragOffset = { x: 0, y: 0 };

// For connector tool
let connectorStartObj = null;

canvas.style.touchAction = "none";
function highlightTool(selectedTool) {
    toolButtons.forEach(t => {
        let btn = document.getElementById("tool-" + t);
        if (btn) {
            btn.style.background = (t === selectedTool) ? "#2563eb" : "#f1f5f9";
            btn.style.color = (t === selectedTool) ? "#fff" : "#222";
            btn.style.transition = "all 0.17s";
        }
    });
}
document.addEventListener('keydown', e => {
    if ((e.ctrlKey || e.metaKey) && e.key === 'z') { undo(); }
    if ((e.ctrlKey || e.metaKey) && e.key === 'y') { redo(); }
    if (e.key === "v") setTool('select');
    if (e.key === "p") setTool('pen');
});
function setTool(t){
    tool = t;
    selected = null;
    dragging = false;
    connectorStartObj = null;
    canvas.style.borderColor = (tool === "select") ? "#039855" : "#2563eb";
    highlightTool(tool);
    render();
}
function isTouchDevice() {
    return 'ontouchstart' in window || navigator.maxTouchPoints > 0;
}
function getCanvasEvent(e) {
    if (e.touches && e.touches.length) return { clientX: e.touches[0].clientX, clientY: e.touches[0].clientY };
    if (e.changedTouches && e.changedTouches.length) return { clientX: e.changedTouches[0].clientX, clientY: e.changedTouches[0].clientY };
    return e;
}
const downEvent = isTouchDevice() ? "touchstart" : "mousedown";
const moveEvent = isTouchDevice() ? "touchmove" : "mousemove";
const upEvent = isTouchDevice() ? "touchend" : "mouseup";
canvas.addEventListener('touchmove', function (e) { e.preventDefault(); }, { passive: false });

canvas.addEventListener(downEvent, (e)=>{
    const pointer = getCanvasEvent(e);
    const {x,y} = getXY(pointer);
    startX = x; startY = y;

    if(tool === "connector") {
        let obj = getHitObject(x, y, {preferCenters: true});
        if(!obj || (obj.type !== "rect" && obj.type !== "circle" && obj.type !== "text")) {
            showMsg("Tap on a rectangle, circle, or text to start or complete a connection.", true);
            return;
        }
        if(!connectorStartObj) {
            connectorStartObj = obj;
            selected = obj;
            showMsg("Now select the next object to connect.", false);
            render();
        } else if (connectorStartObj.id !== obj.id) {
            let from = getObjectCenter(connectorStartObj), to = getObjectCenter(obj);
            if (!connectorStartObj.id) connectorStartObj.id = rid();
            if (!obj.id) obj.id = rid();
            actions.push({
                id: rid(),
                type: "connector",
                x1: from.x, y1: from.y, x2: to.x, y2: to.y,
                fromId: connectorStartObj.id,
                toId: obj.id,
                color: "#039855",
                size: 4
            });
            redoStack = [];
            selected = null;
            connectorStartObj = null;
            render();
            showMsg("Objects connected!", true);
        }
        return;
    }

    if(tool === "select"){
        selected = getHitObject(x,y);
        if(selected){
            dragging = true;
            if(selected.type === "rect" || selected.type === "circle" || selected.type === "text") {
                dragOffset.x = x - selected.x;
                dragOffset.y = y - selected.y;
            } else if(selected.type === "line" || selected.type === "arrow" || selected.type === "connector") {
                dragOffset.x = x - selected.x1;
                dragOffset.y = y - selected.y1;
            } else if(selected.type === "pen" && selected.path.length) {
                dragOffset.x = x - selected.path[0][0];
                dragOffset.y = y - selected.path[0][1];
            }
            showMsg("Drag to move. Click elsewhere to deselect.", true);
        }
        return;
    }
    if(tool === "text"){
        setTimeout(() => { 
            const text = prompt("Enter text to add:");
            if(text){
                actions.push({
                    id: rid(),
                    type: "text",
                    text,
                    x,y,
                    color: getColor()
                });
                render();
                showMsg("Text added ✅", true);
            }
        }, 18);
        return;
    }
    drawing = true;
    if(tool === "pen"){
        currentPath = [[x,y]];
    }
});
canvas.addEventListener(moveEvent,(e)=>{
    if(!(drawing || dragging)) return;
    const pointer = getCanvasEvent(e);
    const {x,y} = getXY(pointer);

    if(dragging && selected){
        moveObject(selected, x - dragOffset.x, y - dragOffset.y);
        if (selected.type === "rect" || selected.type === "circle" || selected.type === "text") {
            actions.forEach(a => {
                if (a.type === "connector") {
                    if (a.fromId === selected.id) {
                        let newC = getObjectCenter(selected);
                        a.x1 = newC.x;
                        a.y1 = newC.y;
                    }
                    if (a.toId === selected.id) {
                        let newC = getObjectCenter(selected);
                        a.x2 = newC.x;
                        a.y2 = newC.y;
                    }
                }
            });
        }
        render();
        return;
    }

    if(tool === "pen" && drawing){
        currentPath.push([x,y]);
        render();
    }
});
canvas.addEventListener(upEvent,(e)=>{
    if(!(drawing || dragging)) return;
    drawing = false;
    dragging = false;
    const pointer = getCanvasEvent(e);
    const {x,y} = getXY(pointer);
    const color = getColor();
    const size = getSize();

    if(tool === "pen" && currentPath.length > 1){
        actions.push({id:rid(), type:"pen", path:currentPath.slice(), color, size});
        showMsg("Stroke added.", true);
    }
    if(tool === "rect" && (x!==startX || y!==startY)){
        actions.push({id:rid(), type:"rect",
            x:startX, y:startY, w:x-startX, h:y-startY, color, size});
        showMsg("Rectangle added.", true);
    }
    if(tool === "circle" && (x!==startX || y!==startY)){
        actions.push({id:rid(), type:"circle",
            x:startX, y:startY, r:dist(startX,startY,x,y), color, size});
        showMsg("Circle added.", true);
    }
    if(tool === "line" && (x!==startX || y!==startY)){
        actions.push({id:rid(), type:"line",
            x1:startX, y1:startY, x2:x, y2:y, color, size});
        showMsg("Line added.", true);
    }
    if(tool === "arrow" && (x!==startX || y!==startY)){
        actions.push({id:rid(), type:"arrow",
            x1:startX, y1:startY, x2:x, y2:y, color, size});
        showMsg("Arrow added.", true);
    }

    redoStack = [];
    render();
});
function moveObject(obj,newX,newY){
    if(obj.type === "rect" || obj.type==="circle" || obj.type==="text"){
        obj.x = newX; obj.y = newY;
    }
    if(obj.type==="line" || obj.type==="arrow" || obj.type === "connector"){
        let dx = newX - obj.x1, dy = newY - obj.y1;
        obj.x1 += dx; obj.y1 += dy;
        obj.x2 += dx; obj.y2 += dy;
    }
    if(obj.type==="pen" && obj.path.length){
        let dx = newX - obj.path[0][0], dy = newY - obj.path[0][1];
        obj.path = obj.path.map(p=>[p[0]+dx,p[1]+dy]);
    }
}
function getHitObject(x,y, opts={}) {
    for(let i=actions.length-1;i>=0;i--){
        let a = actions[i];
        if(opts.preferCenters) {
            if(a.type==="rect" || a.type==="circle" || a.type==="text") {
                let c = getObjectCenter(a);
                if(Math.abs(x-c.x) < 24 && Math.abs(y-c.y) < 24) return a;
            }
        } else {
            if(hit(a,x,y)) return a;
        }
    }
    return null;
}
function hit(a,x,y){
    if(a.type==="rect") return x>a.x && x<a.x+a.w && y>a.y && y<a.y+a.h;
    if(a.type==="circle") return Math.hypot(x-a.x,y-a.y)<a.r;
    if(a.type==="text") return Math.abs(x-a.x)<50 && Math.abs(y-a.y)<20;
    if(a.type==="line"||a.type==="arrow"||a.type==="connector") return pointLineDist(x,y,a)<10;
    if(a.type==="pen") return a.path.some(p=>Math.hypot(x-p[0],y-p[1])<8);
    return false;
}
function pointLineDist(px,py,a){
    let x1=a.x1,y1=a.y1,x2=a.x2,y2=a.y2;
    let A=px-x1,B=py-y1,C=x2-x1,D=y2-y1;
    let dot=A*C+B*D,len=C*C+D*D;
    let t=Math.max(0,Math.min(1,len>0 ? dot/len : 0));
    let xx=x1+t*C,yy=y1+t*D;
    return Math.hypot(px-xx,py-yy);
}
function getObjectCenter(obj) {
    if (obj.type === "rect") {
        if (typeof obj.w === 'undefined' || typeof obj.h === 'undefined') {
            return { x: obj.x, y: obj.y };
        }
        return { x: obj.x + obj.w/2, y: obj.y + obj.h/2 };
    } else if (obj.type === "circle") {
        return { x: obj.x, y: obj.y };
    } else if (obj.type === "text") {
        return { x: obj.x, y: obj.y };
    }
    return { x: 0, y: 0 };
}
function render(){
    ctx.clearRect(0,0,canvas.width,canvas.height);
    actions.forEach(a=>{
        ctx.strokeStyle=a.color||"#111";
        ctx.fillStyle=a.color||"#111";
        ctx.lineWidth=a.size||2;

        if(a.type==="pen"){
            ctx.beginPath();
            a.path.forEach((p,i)=>{
                if(i===0) ctx.moveTo(p[0],p[1]);
                else ctx.lineTo(p[0],p[1]);
            });
            ctx.stroke();
        }
        if(a.type==="rect") ctx.strokeRect(a.x,a.y,a.w,a.h);
        if(a.type==="circle"){
            ctx.beginPath();
            ctx.arc(a.x,a.y,a.r,0,Math.PI*2);
            ctx.stroke();
        }
        if(a.type==="line"){
            ctx.beginPath();
            ctx.moveTo(a.x1,a.y1);
            ctx.lineTo(a.x2,a.y2);
            ctx.stroke();
        }
        if(a.type==="connector"){
            ctx.save();
            ctx.strokeStyle = a.color || "#039855";
            ctx.lineWidth = a.size || 4;
            ctx.beginPath();
            ctx.moveTo(a.x1, a.y1);
            ctx.lineTo(a.x2, a.y2);
            ctx.stroke();
            ctx.restore();
            drawArrowHead(a.x1, a.y1, a.x2, a.y2, a.color || "#039855");
        }
        if(a.type==="arrow") drawArrow(a);
        if(a.type==="text"){
            ctx.font="17px Arial";
            ctx.shadowColor="rgba(0,0,0,0.04)";
            ctx.shadowBlur=8;
            ctx.fillText(a.text,a.x,a.y+6);
            ctx.shadowBlur=0;
        }
        if(selected && selected.id===a.id){
            ctx.save();
            ctx.strokeStyle="#f59e42";
            ctx.lineWidth=2.5;
            if(selected.type === "rect" || selected.type === "circle" || selected.type === "text") {
                ctx.strokeRect(selected.x-10, selected.y-10, 24, 24);
            } else if(selected.type === "line" || selected.type === "arrow" || selected.type === "connector") {
                ctx.strokeRect(selected.x1-10, selected.y1-10, 24, 24);
            } else if(selected.type === "pen" && selected.path.length) {
                ctx.strokeRect(selected.path[0][0]-10, selected.path[0][1]-10, 24, 24);
            }
            ctx.restore();
        }
    });

    if(tool === "connector" && connectorStartObj) {
        ctx.save();
        const c = getObjectCenter(connectorStartObj);
        ctx.strokeStyle = "#f59e42";
        ctx.lineWidth = 2.5;
        ctx.beginPath();
        ctx.arc(c.x, c.y, 18, 0, Math.PI*2);
        ctx.stroke();
        ctx.restore();
    }
}
function drawArrow(a){
    const head=14;
    const angle=Math.atan2(a.y2-a.y1,a.x2-a.x1);
    ctx.beginPath();
    ctx.moveTo(a.x1,a.y1);
    ctx.lineTo(a.x2,a.y2);
    ctx.stroke();

    ctx.beginPath();
    ctx.moveTo(a.x2,a.y2);
    ctx.lineTo(a.x2-head*Math.cos(angle-Math.PI/6), a.y2-head*Math.sin(angle-Math.PI/6));
    ctx.moveTo(a.x2,a.y2);
    ctx.lineTo(a.x2-head*Math.cos(angle+Math.PI/6), a.y2-head*Math.sin(angle+Math.PI/6));
    ctx.stroke();
}
function drawArrowHead(x1, y1, x2, y2, color="#039855") {
    const headLength = 16;
    const angle = Math.atan2(y2-y1, x2-x1);
    ctx.save();
    ctx.strokeStyle = color;
    ctx.lineWidth = 3.8;
    ctx.beginPath();
    ctx.moveTo(x2, y2);
    ctx.lineTo(x2 - headLength * Math.cos(angle - Math.PI/7), y2 - headLength * Math.sin(angle - Math.PI/7));
    ctx.moveTo(x2, y2);
    ctx.lineTo(x2 - headLength * Math.cos(angle + Math.PI/7), y2 - headLength * Math.sin(angle + Math.PI/7));
    ctx.stroke();
    ctx.restore();
}
function undo(){
    if(actions.length){
        redoStack.push(actions.pop());
        render();
        showMsg("Undo.", true);
    }
}
function redo(){
    if(redoStack.length){
        actions.push(redoStack.pop());
        render();
        showMsg("Redo.", true);
    }
}
function clearCanvas(){
    if (actions.length === 0) {
        showMsg("Canvas is already clear.", true);
        return;
    }
    if (confirm('Clear all elements from the canvas?')){
        actions=[]; redoStack=[]; render();
        showMsg("Canvas cleared.", true);
    }
}
function getXY(e){
    const r=canvas.getBoundingClientRect();
    return {x:e.clientX-r.left,y:e.clientY-r.top};
}
function getColor(){
    return document.getElementById("colorPicker").value;
}
function getSize(){
    return +document.getElementById("brushSize").value;
}
function dist(x1,y1,x2,y2){
    return Math.hypot(x2-x1,y2-y1);
}
function rid(){
    return (window.crypto && window.crypto.randomUUID) ? crypto.randomUUID() : (Date.now().toString(36)+Math.random().toString(36).substr(2,5));
}
function showMsg(msg, fade=false){
    let el = document.getElementById("canvas-message");
    el.textContent = msg;
    el.style.display = "block";
    el.style.opacity = "1";
    if(fade) setTimeout(()=>{
        el.style.opacity = '0';
        setTimeout(()=>{el.style.display='none'}, 400);
    }, 1050);
}
function saveCanvas(){
    if (actions.length === 0) {
        showMsg("🖍️ Please draw your workflow before saving!"); 
        let btn = document.getElementById("save-btn");
        if(btn){ btn.disabled = true; setTimeout(()=>{btn.disabled=false}, 1600);}
        return;
    }
    document.getElementById("canvas_data").value = JSON.stringify(actions);
    let btn = document.getElementById("save-btn");
    let msg = document.getElementById("canvas-message");
    btn.disabled = true;
    msg.textContent = "Saving your workflow ...";
    msg.style.display = "block";
    setTimeout(() => {
        document.getElementById("canvasForm").submit();
    }, 400);
}
window.onload = () => { 
    highlightTool(tool); 
    render(); 
};

</script>
'''
    return render_template_string(BASE_HTML, body=body, progress=85)

# =========================================================
# REVIEW
# =========================================================

# WORKFLOW CANVAS - KYC Workflow Detailed Stepwise Explanation with More Features
@app.route('/gap', methods=['GET', 'POST'])
def gap():
    import sqlite3
    from flask import render_template_string, request, redirect, session, flash, url_for
    import html

    if request.method == 'POST':
        explanation = request.form.get('explanation', '')
        heading_color = request.form.get('heading_color', '#222')
        heading_size = request.form.get('heading_size', '18')
        section_title = request.form.get('section_title', 'KYC Workflow Explanation')
        section_icon = request.form.get('section_icon', '📝')
        allow_numbering = request.form.get('allow_numbering', '') == 'on'
        candidate_id = session.get('candidate_id')

        if not candidate_id:
            flash("Your session expired. Please sign in again to continue.", "error")
            return redirect(url_for('gap'))

        if not explanation or explanation.strip() == "":
            flash("📝 Please provide your point-wise/bulleted workflow explanation before submitting.", "warning")
            return redirect(url_for('gap'))

        # Optional: Number items if enabled
        lines = explanation.split('\n')
        if allow_numbering:
            numbered = []
            count = 1
            for line in lines:
                l = line.strip()
                if l.startswith('- '):
                    l = f"{count}. {l[2:]}"
                    count += 1
                numbered.append(l)
            explanation_final = '\n'.join(numbered)
        else:
            explanation_final = explanation

        try:
            conn = sqlite3.connect(DB)
            c = conn.cursor()
            # Save with additional section title and icon support
            formatted_explanation = f"""
                <div>
                    <div style='color:{heading_color}; font-size:{heading_size}px; font-weight:700; margin-bottom:12px;display:flex;gap:7px;align-items:center;'>
                        <span>{section_icon}</span>
                        <span>{html.escape(section_title)}</span>
                    </div>
                    <div style='font-size:16px;color:#233;white-space:pre-line;background:#f8fbff;padding:10px 14px;border-radius:8px;border:1px solid #e0e7ef;margin-bottom:8px;'>
                        {html.escape(explanation_final)}
                    </div>
                </div>
            """
            c.execute("""
                INSERT INTO answers(candidate_id, question, answer)
                VALUES (?, ?, ?)
            """, (candidate_id, section_title, formatted_explanation))
            conn.commit()
        except Exception as e:
            flash("Oops! Something went wrong while saving your workflow. Please try again. " + str(e), "error")
            return redirect(url_for('gap'))
        finally:
            conn.close()
        flash("✅ Workflow explanation saved successfully! Redirecting to next step ...", "success")
        return redirect('/review')

    # Enhanced: custom section title, icon, toggle numbered list, copy/paste, preview
    body = '''
    <h2 style="margin-bottom: 12px;">KYC Workflow Stepwise Explanation</h2>
    <div class="question-box" style="padding-bottom:24px;">
      <div style="background: #e3f3ff; border-radius: 10px; margin-bottom: 18px; padding: 14px 22px; color: #264366;box-shadow:0 2px 14px #3476e014">
        <span style="font-size:18px;font-weight:500;">📝 &nbsp;Describe your KYC workflow step by step, with features such as bullets, numbering, section title/icon, preview, and copy/paste.</span><br>
        <span style="color:#5d6887;">
        <b>Note:</b> Use ➕ Bullet or #️⃣ Numbered to insert bullets or automatic numbering.<br>
        <b>You may also add a custom heading/title, select a color and size, and pick an icon for your section heading. Click 👁️ Preview to see your formatted output before saving.<br>
        <b>Tip:</b> Select text below and press "Copy" to quickly keep a local backup.<br>
        <br>
        <i>Example:
        <div style="margin-top:4px;color:#222;padding-left:14px">
        - Step 1: Collect customer documents <br>
        - Step 2: Verify identity and address <br>
        - Step 3: Perform risk assessment <br>
        - Step 4: Approve or reject application
        </div>
        </i>
        </span>
      </div>
      <form method="POST" id="canvasForm" style="margin-top:12px;text-align:right;" autocomplete="off">
        <div style="margin-bottom:18px; text-align:left; max-width:760px;">
            <label for="section_title" style="font-weight:600;color:#25506e;font-size:15px;display:block;margin-bottom:5px;">
                Heading/Section Title:
                <input type="text" name="section_title" id="section_title" maxlength="80" value="KYC Workflow Explanation" style="padding:4px 9px; border-radius:7px; border:1.1px solid #accbfc; font-size:15px; width:60%; margin-left:10px;">
            </label>
            <div style="margin-bottom:7px;display:flex;align-items:center;gap:18px;">
                <span style="font-size:15px;">
                  Icon: 
                  <select name="section_icon" id="section_icon" style="font-size:18px;margin-left:2px;">
                    <option value="📝" selected>📝</option>
                    <option value="✔️">✔️</option>
                    <option value="✅">✅</option>
                    <option value="🔎">🔎</option>
                    <option value="🔐">🔐</option>
                    <option value="⭐">⭐</option>
                    <option value="🔶">🔶</option>
                    <option value="🛡️">🛡️</option>
                    <option value="➡️">➡️</option>
                  </select>
                </span>
                <span style="font-size:15px;">
                  Heading Color: <input type="color" name="heading_color" id="heading_color" value="#2563eb" style="margin-left:2px;">
                </span>
                <span style="font-size:15px;">
                  Heading Size:
                  <select name="heading_size" id="heading_size" style="margin-left:2px;">
                    <option value="17">17px</option>
                    <option value="18" selected>18px</option>
                    <option value="20">20px</option>
                    <option value="22">22px</option>
                    <option value="24">24px</option>
                  </select>
                </span>
                <span style="font-size:15px;"><input type="checkbox" name="allow_numbering" id="allow_numbering" style="margin-right:4px;">#️⃣ Numbered</span>
                <button type="button" class="btn" style="font-size:15px;" onclick="previewWorkflow()" title="Preview your formatted workflow">👁️ Preview</button>
                <button type="button" class="btn" style="font-size:15px;" onclick="copyToClipboard()" title="Copy explanation text">📋 Copy</button>
            </div>
            <label for="explanation" style="font-weight:600;color:#29405e;display:block;font-size:16px;margin-bottom:7px;">
                <span style="font-size:17px;">📝</span> Please write your KYC workflow below (point-wise/bullets or numbered, no diagram):
            </label>
            <div style="margin-top:7px; display:flex; align-items:center; gap:16px;margin-bottom:3px;">
                <button type="button" class="btn" onclick="insertBulletPoint()" title="Insert Bullet Point">➕ Bullet</button>
                <button type="button" class="btn" onclick="insertNumberedPoint()" title="Insert Numbered Point">#️⃣ Numbered</button>
                <span style="font-size:14px;color:#6b7280;">(Use Enter to add more points)</span>
            </div>
            <textarea name="explanation" id="explanation" rows="8"
                style="width:100%;margin-top:5px;padding:9px 13px;border-radius:8px;border:1.5px solid #c3dbea;font-size:15.7px;resize:vertical;min-height:110px;"
                placeholder="e.g.\n- Step 1: Collect customer documents\n- Step 2: Verify identity and address\n- Step 3: Perform risk assessment\n- Step 4: Approve or reject application"></textarea>
            <div style="color:#64748b;font-size:13px;margin-top:3px;">
                Only point-wise explanation is required. Please do not draw a canvas/diagram for this step. (Formatting tools available above.)
            </div>
            <div id="preview-area" style="background:#f3f6fb;margin-top:14px;padding:10px 14px;border-radius:7px;min-height:32px; font-size:15.5px; display:none;">
                <!-- Preview appears here -->
            </div>
        </div>
        <button type="submit" class="btn" id="save-btn" style="font-weight:700;font-size:17px;margin-top:12px;">
          💾 Save & Continue
        </button>
      </form>
      <div id="canvas-message" style="margin-top:18px;color:#2563eb;font-size:15px;display:none;transition:opacity 0.2s;"></div>
    </div>
<script>
// Insert Bullet
function insertBulletPoint() {
    var textarea = document.getElementById("explanation");
    if (!textarea) return;
    var start = textarea.selectionStart, value = textarea.value;
    var lineStart = value.lastIndexOf('\\n', start - 1) + 1;
    var insert = "- ";
    if (value.substr(lineStart, 2) === "- ") {
        textarea.selectionStart = textarea.selectionEnd = lineStart + 2; textarea.focus(); return;
    }
    textarea.value = value.slice(0, lineStart) + insert + value.slice(lineStart);
    var newPos = start + insert.length;
    textarea.selectionStart = textarea.selectionEnd = newPos;
    textarea.focus();
}
// Insert Numbered Point
function insertNumberedPoint() {
    var textarea = document.getElementById("explanation");
    if (!textarea) return;
    var value = textarea.value;
    var start = textarea.selectionStart;
    var lineStart = value.lastIndexOf('\\n', start-1) + 1;
    // Find previous line with numbered pattern
    var prevNumber = 0;
    var prevLines = value.slice(0, lineStart).split('\\n');
    for(var i=prevLines.length-1;i>=0;i--) {
        var match = prevLines[i].match(/^\\s*(\\d+)\\. /);
        if(match){prevNumber = parseInt(match[1]);break;}
    }
    var insert = (prevNumber+1) + ". ";
    textarea.value = value.slice(0, lineStart) + insert + value.slice(lineStart);
    var newPos = start + insert.length;
    textarea.selectionStart = textarea.selectionEnd = newPos;
    textarea.focus();
}
// Preview Section
function previewWorkflow() {
    var heading = document.getElementById("section_title").value.trim();
    var icon = document.getElementById("section_icon").value;
    var color = document.getElementById("heading_color").value;
    var size = document.getElementById("heading_size").value;
    var txt = document.getElementById("explanation").value;
    var numbered = document.getElementById("allow_numbering").checked;
    var out = "";
    // Escape HTML
    function esc(x) {
        return x.replace(/</g,"&lt;").replace(/>/g,"&gt;");
    }
    if (heading) {
        out += "<div style='color:"+color+";font-size:"+size+"px;font-weight:700;margin-bottom:12px;display:flex;gap:7px;align-items:center;'><span>"+icon+"</span> <span>"+esc(heading)+"</span></div>";
    }
    // Format body as numbered or bullets, preserving lines
    var lines = txt.trim().split(/\\r?\\n/);
    if (numbered) {
        let num = 1;
        lines = lines.map(l => {
            if(l.trim().startsWith('- ')) {
                return (num++) + '. ' + esc(l.trim().slice(2));
            } else if(l.trim().match(/^\\d+\\.\\s/)) {
                return esc(l.trim());
            } else if(l.trim()!=='') {
                return (num++) + '. ' + esc(l.trim());
            } else { return ''; }
        });
    } else {
        lines = lines.map(l => esc(l));
    }
    out += "<div style='font-size:16px;color:#233;white-space:pre-line;background:#f8fbff;padding:10px 14px;border-radius:8px;border:1px solid #e0e7ef;margin-bottom:8px;'>" +
        lines.join('\\n') +
        "</div>";
    var previewArea = document.getElementById("preview-area");
    previewArea.innerHTML = out;
    previewArea.style.display = "block";
    previewArea.scrollIntoView({behavior:"smooth"});
}
// Copy Text to Clipboard
function copyToClipboard() {
    var textarea = document.getElementById("explanation");
    if (!textarea) return;
    textarea.select();
    document.execCommand("copy");
    var msg = document.getElementById("canvas-message");
    msg.textContent = "Copied to clipboard!";
    msg.style.display = "block";
    msg.style.opacity = "1";
    setTimeout(()=>{ msg.style.opacity = '0'; setTimeout(()=>{msg.style.display='none'}, 400); }, 900);
}
</script>
'''
    return render_template_string(BASE_HTML, body=body, progress=85)


@app.route('/review')
def review():
    conn = sqlite3.connect(DB)
    c = conn.cursor()

    # Fetch candidate details: name, email, phone
    candidate_id = session.get('candidate_id')
    c.execute('SELECT name, email, phone FROM candidates WHERE id=?', (candidate_id,))
    details_row = c.fetchone()
    if details_row:
        candidate_name = details_row[0] or "Unknown Candidate"
        candidate_email = details_row[1] or "N/A"
        candidate_phone = details_row[2] or "N/A"
    else:
        candidate_name = "Unknown Candidate"
        candidate_email = "N/A"
        candidate_phone = "N/A"

    c.execute('''
        SELECT question, answer
        FROM answers
        WHERE candidate_id=?
    ''', (
        candidate_id,
    ))
    rows = c.fetchall()
    conn.close()
    body = f'<h2>Final Review</h2>'
    body += f'''
    <div style="margin:12px 0 30px;font-size:20px;font-weight:500;color:#226;">
        👤 Candidate: <span style="background:#e8f3ff;padding:3px 12px;border-radius:8px;">{candidate_name}</span><br>
        <span style="font-size:16px;font-weight:400;color:#254b77;">
            📧 Email: <span style="background:#f8fbff;padding:2px 11px;border-radius:8px;">{candidate_email}</span><br>
            📱 Phone: <span style="background:#f5faff;padding:2px 11px;border-radius:8px;">{candidate_phone}</span>
        </span>
    </div>'''
    for r in rows:
        body += f'''
        <div class="question-box">
            <div class="question">
            {r[0]}</div>
            <div style="margin-top:4px;margin-left:1px;">{r[1]}</div>
        </div>
        '''
    body += '''
    <a href="/submit">
      <button class="btn" style="float:right;margin-top:12px;">
      ✅ Final Submit
      </button>
    </a>
    '''
    return render_template_string(
        BASE_HTML,
        body=body,
        progress=100
    )

# =========================================================
# SUBMIT
# =========================================================

@app.route('/submit')
def submit():
    import resend
    import base64
    import os
    import tempfile

    # Submit the exam (update db)
    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute('''
        UPDATE candidates
        SET submit_time=?
        WHERE id=?
    ''', (
        str(datetime.now()),
        session['candidate_id']
    ))
    conn.commit()
    
    # Gather candidate info
    c.execute('SELECT name, email, phone FROM candidates WHERE id=?', (session['candidate_id'],))
    row = c.fetchone()
    candidate = {
        'name': row[0] if row else 'Unknown',
        'email': row[1] if row else 'N/A',
        'phone': row[2] if row else 'N/A',
    }

    # Gather answers
    c.execute('SELECT question, answer FROM answers WHERE candidate_id=?', (session['candidate_id'],))
    answers = c.fetchall()
    conn.close()
    
    # Prepare temp text report
    report_lines = [
        f"Candidate Name: {candidate['name']}",
        f"Email: {candidate['email']}",
        f"Phone: {candidate['phone']}",
        "\n--- Answer Sheet ---"
    ]
    for i, (q, a) in enumerate(answers, 1):
        report_lines.append(f"\nQ{i}: {q}\nA{i}: {a}")

    # Use the system temp directory that works cross-platform
    with tempfile.NamedTemporaryFile(mode='w+', delete=False, suffix='_exam_report.txt', prefix=f"{candidate['name'].replace(' ', '_')}_", encoding='utf-8') as tmp_file:
        tmp_file.write('\n'.join(report_lines))
        temp_txt_path = tmp_file.name

    try:
        resend.api_key = "re_hW5ToeJj_KzXHNfzwsgFVZMZ4f5EYy8NR"  # actual API key
        
        # Read the file as binary for proper attachment encoding
        with open(temp_txt_path, "rb") as f:
            file_bytes = f.read()
        
        file_bytes_b64 = base64.b64encode(file_bytes).decode("utf-8")
        file_payload = {
            "content": file_bytes_b64,
            "filename": os.path.basename(temp_txt_path),
            "type": "text/plain"
        }
        
        text_body = (
            f"Dear HR,\n\n"
            f"Please find attached the detailed submission report for candidate: {candidate['name']}.\n\n"
            "Regards,\nExam System"
        )
        
        resend.Emails.send({
            "from": "onboarding@resend.dev",
            "to": ["roopamugappa@gmail.com"],
            "subject": f"Python Exam Submission: {candidate['name']}",
            "text": text_body,
            "attachments": [file_payload]
        })
    except Exception as e:
        print(f"Error sending email via resend: {e}")

    # Remove temp file if present
    try:
        if os.path.exists(temp_txt_path):
            os.remove(temp_txt_path)
    except Exception as e:
        print(f"Error deleting temp report: {e}")

    # Render success page
    body = '''
    <h2 style="font-size:2.2em;text-align:center;color:var(--success);margin-top:40px;">
    🎉 Assessment Submitted Successfully!
    </h2>
    <script>localStorage.removeItem("timeLeft");</script>
    <div style="max-width:380px;margin:44px auto 0;font-size:17px;">Thank you for completing the assessment.<br/>You may close this tab or window.</div>
    '''
    return render_template_string(
        BASE_HTML,
        body=body,
        progress=100
    )

# =========================================================
# RUN
# =========================================================

if __name__ == '__main__':
    app.run(debug=True, port=5001)


